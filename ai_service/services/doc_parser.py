"""
文档解析服务，支持多种格式的文档内容提取
"""
import os
from typing import List, Dict
import pdfplumber
from docx import Document
import pandas as pd
import markdown
from bs4 import BeautifulSoup
import json
from utils.logger import doc_parser_logger as logger
import tempfile
import moviepy.editor as mp
import whisper
import pytesseract
from PIL import Image
pytesseract.pytesseract.tesseract_cmd = r"D:\Tesseract-OCR\tesseract.exe"

class DocumentParser:
    """文档解析器，支持多种文档格式的内容提取"""
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.txt', '.md', '.json', '.csv', '.xlsx', '.html', '.mp4', '.avi', '.mov', '.mkv', '.jpg', '.jpeg', '.png', '.bmp', '.tiff'}
    
    @staticmethod
    def parse_pdf(file_path: str) -> List[Dict[str, str]]:
        """解析PDF文件，支持图片型PDF OCR"""
        try:
            chunks = []
            with pdfplumber.open(file_path) as pdf:
                empty_pages = 0
                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text()
                    if not text:
                        # 尝试OCR图片型PDF
                        pil_img = page.to_image(resolution=300).original
                        ocr_text = pytesseract.image_to_string(pil_img, lang='chi_sim+eng')
                        if ocr_text.strip():
                            paragraphs = ocr_text.split('\n\n')
                            for para in paragraphs:
                                if para.strip():
                                    chunks.append({
                                        'content': para.strip(),
                                        'source': f"{os.path.basename(file_path)} 第{page_num}页(OCR)"
                                    })
                        else:
                            logger.warning(f"PDF {file_path} 第{page_num}页无可提取文本（文本+OCR均无）")
                            empty_pages += 1
                    else:
                        paragraphs = text.split('\n\n')
                        for para in paragraphs:
                            if para.strip():
                                chunks.append({
                                    'content': para.strip(),
                                    'source': f"{os.path.basename(file_path)} 第{page_num}页"
                                })
            if not chunks:
                logger.error(f"PDF {file_path} 无可提取文本，可能是无文本或加密PDF")
                raise ValueError("PDF无可提取文本，可能是无文本图片型PDF或加密PDF")
            logger.info(f"Successfully parsed PDF file: {file_path}，共{len(chunks)}个分块，{empty_pages}页无文本")
            return chunks
        except Exception as e:
            logger.error(f"Error parsing PDF file {file_path}: {str(e)}")
            raise

    @staticmethod
    def parse_docx(file_path: str) -> List[Dict[str, str]]:
        """解析Word文档"""
        try:
            chunks = []
            doc = Document(file_path)
            
            # 处理段落
            for para_num, para in enumerate(doc.paragraphs, 1):
                if para.text.strip():
                    chunks.append({
                        'content': para.text.strip(),
                        'source': f"{os.path.basename(file_path)} 第{para_num}段"
                    })
            
            # 处理表格
            for table_num, table in enumerate(doc.tables, 1):
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        chunks.append({
                            'content': row_text,
                            'source': f"{os.path.basename(file_path)} 表格{table_num}"
                        })
            
            logger.info(f"Successfully parsed DOCX file: {file_path}")
            return chunks
        except Exception as e:
            logger.error(f"Error parsing DOCX file {file_path}: {str(e)}")
            raise

    @staticmethod
    def parse_txt(file_path: str) -> List[Dict[str, str]]:
        """解析文本文件"""
        try:
            chunks = []
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
                paragraphs = text.split('\n\n')
                for para_num, para in enumerate(paragraphs, 1):
                    if para.strip():
                        chunks.append({
                            'content': para.strip(),
                            'source': f"{os.path.basename(file_path)} 第{para_num}段"
                        })
            logger.info(f"Successfully parsed TXT file: {file_path}")
            return chunks
        except Exception as e:
            logger.error(f"Error parsing TXT file {file_path}: {str(e)}")
            raise

    @staticmethod
    def parse_markdown(file_path: str) -> List[Dict[str, str]]:
        """解析Markdown文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                md_text = f.read()
            
            # 转换为HTML
            html = markdown.markdown(md_text)
            soup = BeautifulSoup(html, 'html.parser')
            
            chunks = []
            # 处理标题和段落
            for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p']):
                if element.text.strip():
                    chunks.append({
                        'content': element.text.strip(),
                        'source': f"{os.path.basename(file_path)} {element.name}"
                    })
            
            logger.info(f"Successfully parsed Markdown file: {file_path}")
            return chunks
        except Exception as e:
            logger.error(f"Error parsing Markdown file {file_path}: {str(e)}")
            raise

    @staticmethod
    def parse_json(file_path: str) -> List[Dict[str, str]]:
        """解析JSON文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            chunks = []
            def process_json(obj, path="root"):
                if isinstance(obj, dict):
                    for key, value in obj.items():
                        if isinstance(value, (str, int, float, bool)):
                            chunks.append({
                                'content': f"{key}: {value}",
                                'source': f"{os.path.basename(file_path)} {path}"
                            })
                        else:
                            process_json(value, f"{path}.{key}")
                elif isinstance(obj, list):
                    for i, item in enumerate(obj):
                        process_json(item, f"{path}[{i}]")
            
            process_json(data)
            logger.info(f"Successfully parsed JSON file: {file_path}")
            return chunks
        except Exception as e:
            logger.error(f"Error parsing JSON file {file_path}: {str(e)}")
            raise

    @staticmethod
    def parse_excel(file_path: str) -> List[Dict[str, str]]:
        """解析Excel文件"""
        try:
            chunks = []
            # 读取所有sheet
            excel_file = pd.ExcelFile(file_path)
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                # 将每行转换为文本
                for idx, row in df.iterrows():
                    row_text = ' | '.join(f"{col}: {val}" for col, val in row.items() if pd.notna(val))
                    if row_text:
                        chunks.append({
                            'content': row_text,
                            'source': f"{os.path.basename(file_path)} {sheet_name} 第{idx+1}行"
                        })
            
            logger.info(f"Successfully parsed Excel file: {file_path}")
            return chunks
        except Exception as e:
            logger.error(f"Error parsing Excel file {file_path}: {str(e)}")
            raise

    @staticmethod
    def parse_csv(file_path: str) -> List[Dict[str, str]]:
        """解析CSV文件"""
        try:
            chunks = []
            df = pd.read_csv(file_path)
            
            # 将每行转换为文本
            for idx, row in df.iterrows():
                row_text = ' | '.join(f"{col}: {val}" for col, val in row.items() if pd.notna(val))
                if row_text:
                    chunks.append({
                        'content': row_text,
                        'source': f"{os.path.basename(file_path)} 第{idx+1}行"
                    })
            
            logger.info(f"Successfully parsed CSV file: {file_path}")
            return chunks
        except Exception as e:
            logger.error(f"Error parsing CSV file {file_path}: {str(e)}")
            raise

    @staticmethod
    def parse_html(file_path: str) -> List[Dict[str, str]]:
        """解析HTML文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                soup = BeautifulSoup(f.read(), 'html.parser')
            
            chunks = []
            # 提取文本内容
            for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'div']):
                if element.text.strip():
                    chunks.append({
                        'content': element.text.strip(),
                        'source': f"{os.path.basename(file_path)} {element.name}"
                    })
            
            logger.info(f"Successfully parsed HTML file: {file_path}")
            return chunks
        except Exception as e:
            logger.error(f"Error parsing HTML file {file_path}: {str(e)}")
            raise

    @staticmethod
    def parse_video(file_path: str) -> List[Dict[str, str]]:
        """解析视频文件，提取音频并转写为文本"""
        import shutil

        try:
            # 1. 创建临时WAV文件用于保存音频
            with tempfile.NamedTemporaryFile(suffix='.wav', delete=False) as tmp_audio:
                audio_path = tmp_audio.name
            
            try:
                # 2. 提取音频（使用 moviepy 调用 ffmpeg）
                video = mp.VideoFileClip(file_path)
                if video.audio is None:
                    raise ValueError("视频文件中不包含音频轨道")
                
                # 输出为16KHz、16-bit的WAV格式（Whisper推荐）
                video.audio.write_audiofile(
                    audio_path,
                    fps=16000,            # Whisper推荐采样率
                    codec='pcm_s16le',    # 无损WAV编码
                    verbose=False,
                    logger=None
                )
            except Exception as e:
                logger.error(f"音频提取失败: {file_path}, 错误: {str(e)}")
                raise ValueError(f"音频提取失败: {str(e)}")

            # 3. 使用 whisper 模型进行语音识别
            try:
                model = whisper.load_model("base")
                result = model.transcribe(audio_path, language='zh')
                text = result.get('text', '').strip()
                segments = result.get('segments', [])

                if not text:
                    logger.warning(f"Whisper未识别到音频内容: {file_path}")
                    raise ValueError("Whisper未识别到音频内容")

                chunks = []
                for seg in segments:
                    seg_text = seg.get('text', '').strip()
                    if seg_text:
                        start = seg.get('start', 0)
                        end = seg.get('end', 0)
                        chunks.append({
                            'content': seg_text,
                            'source': f"{os.path.basename(file_path)} {start:.1f}-{end:.1f}s"
                        })

                if not chunks:
                    chunks.append({
                        'content': text,
                        'source': f"{os.path.basename(file_path)} 全部音频"
                    })

                logger.info(f"成功解析视频: {file_path}，分段数: {len(chunks)}")
                return chunks

            except Exception as e:
                logger.error(f"Whisper语音识别失败: {file_path}, 错误: {str(e)}")
                raise

            finally:
                if os.path.exists(audio_path):
                    os.remove(audio_path)  # 删除临时音频文件

        except Exception as e:
            logger.error(f"parse_video 出错: {file_path}, 错误: {str(e)}")
            raise


    @staticmethod
    def parse_image(file_path: str) -> List[Dict[str, str]]:
        """解析图片文件，OCR识别文本"""
        try:
            img = Image.open(file_path)
            text = pytesseract.image_to_string(img, lang='chi_sim+eng')
            if not text.strip():
                logger.error(f"图片 {file_path} OCR无可识别文本")
                raise ValueError("图片OCR无可识别文本")
            # 分段
            chunks = []
            paragraphs = text.split('\n\n')
            for i, para in enumerate(paragraphs, 1):
                if para.strip():
                    chunks.append({
                        'content': para.strip(),
                        'source': f"{os.path.basename(file_path)} 第{i}段"
                    })
            logger.info(f"Successfully parsed image file: {file_path}，共{len(chunks)}个分块")
            return chunks
        except Exception as e:
            logger.error(f"Error parsing image file {file_path}: {str(e)}")
            raise

    def parse_file(self, file_path: str) -> List[Dict[str, str]]:
        """根据文件类型调用对应的解析方法"""
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext not in self.SUPPORTED_EXTENSIONS:
            error_msg = f"Unsupported file type: {ext}"
            logger.error(error_msg)
            raise ValueError(error_msg)
        
        try:
            if ext == '.pdf':
                return self.parse_pdf(file_path)
            elif ext == '.docx':
                return self.parse_docx(file_path)
            elif ext == '.txt':
                return self.parse_txt(file_path)
            elif ext == '.md':
                return self.parse_markdown(file_path)
            elif ext == '.json':
                return self.parse_json(file_path)
            elif ext in ['.xlsx', '.xls']:
                return self.parse_excel(file_path)
            elif ext == '.csv':
                return self.parse_csv(file_path)
            elif ext == '.html':
                return self.parse_html(file_path)
            elif ext in ['.mp4', '.avi', '.mov', '.mkv']:
                return self.parse_video(file_path)
            elif ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff']:
                return self.parse_image(file_path)
        except Exception as e:
            logger.error(f"Error parsing file {file_path}: {str(e)}")
            raise

    def chunk_text(self, text: str, chunk_size: int = 500) -> List[Dict[str, str]]:
        """将长文本分块"""
        try:
            chunks = []
            words = text.split()
            current_chunk = []
            current_size = 0
            
            for word in words:
                current_size += len(word) + 1  # +1 for space
                if current_size > chunk_size:
                    if current_chunk:
                        chunks.append({
                            'content': ' '.join(current_chunk),
                            'source': 'Text Input'
                        })
                    current_chunk = [word]
                    current_size = len(word)
                else:
                    current_chunk.append(word)
                    
            if current_chunk:
                chunks.append({
                    'content': ' '.join(current_chunk),
                    'source': 'Text Input'
                })
            
            logger.info(f"Successfully chunked text into {len(chunks)} chunks")
            return chunks
        except Exception as e:
            logger.error(f"Error chunking text: {str(e)}")
            raise 
